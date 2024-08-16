from docx import Document
import os
import matplotlib.pyplot as plt


def generate_report(csv_path, docx_path):
    print("Generating report...")

    # Create a new DOCX file for the report
    report = Document()

    # Add a title
    report.add_heading('Data Cleaning Report', 0)

    # Summary of the cleaned CSV file
    report.add_heading('CSV File Summary', level=1)
    report.add_paragraph(f"Cleaned CSV file: {os.path.basename(csv_path)}")

    # Add a placeholder for charts
    report.add_paragraph("Charts:")

    # Generate and save a chart
    chart_path = 'chart.png'
    plt.figure(figsize=(6, 4))
    plt.plot([1, 2, 3], [4, 5, 6])
    plt.title("Sample Chart")
    plt.savefig(chart_path)

    report.add_picture(chart_path, width=4000000, height=3000000)

    # Summary of the cleaned DOCX file
    report.add_heading('DOCX File Summary', level=1)
    report.add_paragraph(f"Cleaned DOCX file: {os.path.basename(docx_path)}")

    # Save the report
    report_path = 'cleaning_report.docx'
    report.save(report_path)

    print(f"Report generated and saved as: {report_path}")
    return report_path
